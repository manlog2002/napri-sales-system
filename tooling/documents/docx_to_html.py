from pathlib import Path
from base64 import b64encode
from html import escape
from docx import Document
from docx.oxml.ns import qn


ROOT = Path(r"C:\Users\DELL\Documents\Codex\2026-08-19\referenced-chatgpt-conversation-this-is-an")
OUT = ROOT / "outputs" / "حزمة-نبري-البيعية"
HTML_DIR = ROOT / "work" / "napri-sales-kit" / "html"
HTML_DIR.mkdir(parents=True, exist_ok=True)


def paragraph_alignment(p):
    val = str(p.alignment)
    if "CENTER" in val:
        return "center"
    if "LEFT" in val:
        return "left"
    if "JUSTIFY" in val:
        return "justify"
    return "right"


def num_formats(doc):
    numbering = doc.part.numbering_part.element
    abstract_fmt = {}
    for abstract in numbering.findall(qn("w:abstractNum")):
        aid = abstract.get(qn("w:abstractNumId"))
        fmt = abstract.find(".//" + qn("w:numFmt"))
        abstract_fmt[aid] = fmt.get(qn("w:val")) if fmt is not None else "bullet"
    mapping = {}
    for num in numbering.findall(qn("w:num")):
        nid = num.get(qn("w:numId"))
        aid = num.find(qn("w:abstractNumId"))
        mapping[nid] = abstract_fmt.get(aid.get(qn("w:val")) if aid is not None else "", "bullet")
    return mapping


def image_html(doc, paragraph):
    out = []
    for blip in paragraph._p.iter(qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if not rid:
            continue
        part = doc.part.related_parts.get(rid)
        if not part:
            continue
        content_type = getattr(part, "content_type", "image/png")
        data = b64encode(part.blob).decode("ascii")
        out.append(f'<img src="data:{content_type};base64,{data}" alt="لقطة من نظام نبري"/>')
    return "".join(out)


def runs_html(paragraph):
    chunks = []
    for run in paragraph.runs:
        text = escape(run.text).replace("\n", "<br>")
        if not text:
            continue
        if run.bold:
            text = f"<strong>{text}</strong>"
        if run.italic:
            text = f"<em>{text}</em>"
        color = None
        if run.font.color and run.font.color.rgb:
            color = str(run.font.color.rgb)
        if color:
            text = f'<span style="color:#{color}">{text}</span>'
        chunks.append(text)
    return "".join(chunks)


def has_page_break(paragraph):
    for br in paragraph._p.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def paragraph_html(doc, paragraph, formats, in_cell=False):
    style = paragraph.style.name if paragraph.style else "Normal"
    align = paragraph_alignment(paragraph)
    images = image_html(doc, paragraph)
    text = runs_html(paragraph)
    num_pr = paragraph._p.pPr.numPr if paragraph._p.pPr is not None else None
    class_names = ["para"]
    tag = "p"
    if style.startswith("Heading 1"):
        tag, class_names = "h1", ["h1"]
    elif style.startswith("Heading 2"):
        tag, class_names = "h2", ["h2"]
    elif style.startswith("Heading 3"):
        tag, class_names = "h3", ["h3"]
    if num_pr is not None and num_pr.numId is not None:
        fmt = formats.get(num_pr.numId.val, "bullet")
        marker = "1." if fmt == "decimal" else "•"
        class_names.append("list-item")
        text = f'<span class="marker">{marker}</span><span>{text}</span>'
    if images:
        text = images + text
        class_names.append("image-para")
    if not text and not has_page_break(paragraph):
        return '<div class="spacer"></div>' if not in_cell else ""
    block = f'<{tag} class="{" ".join(class_names)}" style="text-align:{align}">{text}</{tag}>'
    if has_page_break(paragraph):
        block += '<div class="pagebreak"></div>'
    return block


def cell_fill(cell):
    tc_pr = cell._tc.tcPr
    shd = tc_pr.find(qn("w:shd")) if tc_pr is not None else None
    return shd.get(qn("w:fill")) if shd is not None else None


def table_html(doc, table, formats):
    rows = []
    for row in table.rows:
        seen = set()
        cells_html = []
        for cell in row.cells:
            key = id(cell._tc)
            if key in seen:
                continue
            seen.add(key)
            grid_span = cell._tc.tcPr.find(qn("w:gridSpan")) if cell._tc.tcPr is not None else None
            colspan = int(grid_span.get(qn("w:val"))) if grid_span is not None else 1
            fill = cell_fill(cell)
            style = f' style="background:#{fill}"' if fill and fill not in ("auto", "FFFFFF") else ""
            colspan_attr = f' colspan="{colspan}"' if colspan > 1 else ""
            content = "".join(paragraph_html(doc, p, formats, in_cell=True) for p in cell.paragraphs)
            cells_html.append(f"<td{colspan_attr}{style}>{content}</td>")
        rows.append("<tr>" + "".join(cells_html) + "</tr>")
    return '<table class="doc-table">' + "".join(rows) + "</table>"


def build_html(docx_path):
    doc = Document(docx_path)
    formats = num_formats(doc)
    body = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = next((p for p in doc.paragraphs if p._p is child), None)
            if p is not None:
                body.append(paragraph_html(doc, p, formats))
        elif child.tag == qn("w:tbl"):
            table = next((t for t in doc.tables if t._tbl is child), None)
            if table is not None:
                body.append(table_html(doc, table, formats))
    title = escape(docx_path.stem)
    css = r'''
      @page { size: Letter; margin: 18mm 18mm 18mm 18mm; }
      * { box-sizing: border-box; }
      html, body { margin: 0; padding: 0; background: #fff; color: #14231c; }
      body { font-family: Arial, "Segoe UI", sans-serif; direction: rtl; font-size: 11pt; line-height: 1.55; }
      .para { margin: 0 0 7pt 0; }
      .spacer { height: 11pt; }
      h1 { color: #123c2d; font-size: 16pt; margin: 18pt 0 10pt; page-break-after: avoid; }
      h2 { color: #1d5a43; font-size: 13pt; margin: 14pt 0 7pt; page-break-after: avoid; }
      h3 { color: #14231c; font-size: 12pt; margin: 10pt 0 5pt; page-break-after: avoid; }
      strong { font-weight: 700; }
      .list-item { display: flex; flex-direction: row; gap: 8pt; padding-right: 12pt; margin-bottom: 4pt; }
      .list-item .marker { color: #123c2d; font-weight: 700; min-width: 16pt; text-align: center; }
      .doc-table { width: 100%; border-collapse: collapse; table-layout: fixed; margin: 8pt 0 12pt; break-inside: avoid; }
      .doc-table td { border: 0.6pt solid #d7ddd9; padding: 7pt 8pt; vertical-align: middle; text-align: right; }
      .doc-table p, .doc-table h1, .doc-table h2, .doc-table h3 { margin: 0; }
      img { display: block; max-width: 100%; width: 100%; height: auto; margin: 8pt auto 5pt; border: 0.5pt solid #d7ddd9; border-radius: 6pt; }
      .image-para { break-inside: avoid; }
      .pagebreak { break-after: page; page-break-after: always; }
      a { color: #123c2d; }
    '''
    html = f'<!doctype html><html lang="ar" dir="rtl"><head><meta charset="utf-8"><title>{title}</title><style>{css}</style></head><body>{"".join(body)}</body></html>'
    out_path = HTML_DIR / (docx_path.stem + ".html")
    out_path.write_text(html, encoding="utf-8")
    return out_path


def main():
    for path in sorted(OUT.glob("*.docx")):
        build_html(path)


if __name__ == "__main__":
    main()
