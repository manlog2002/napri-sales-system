from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK


ROOT = Path(r"C:\Users\DELL\Documents\Codex\2026-08-19\referenced-chatgpt-conversation-this-is-an")
WORK = ROOT / "work" / "napri-sales-kit"
ASSETS = WORK / "assets"
OUT = ROOT / "outputs" / "حزمة-نبري-البيعية"
OUT.mkdir(parents=True, exist_ok=True)

GREEN = "123C2D"
GREEN2 = "1D5A43"
LIME = "D9EF9F"
LIME2 = "EEF7D7"
GOLD = "C7963B"
INK = "14231C"
GRAY = "66716B"
LIGHT = "F7F4EC"
WHITE = "FFFFFF"
RED = "A14D45"
BLUEGRAY = "E8EEF5"
VERY_LIGHT = "F2F4F7"


PRESETS = {
    "narrative_proposal": {
        "body_after": 8, "line": 1.333, "align": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "h1_before": 18, "h1_after": 10, "h2_before": 12, "h2_after": 6,
        "h3_before": 8, "h3_after": 4, "table_fill": "F4F6F9",
        "list_left": 0.181, "list_text": 0.375, "list_hanging": 0.194,
        "list_after": 4, "list_line": 1.208,
    },
    "contract_negotiation_brief": {
        "body_after": 6, "line": 1.25, "align": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "h1_before": 14, "h1_after": 8, "h2_before": 11, "h2_after": 6,
        "h3_before": 8, "h3_after": 4, "table_fill": BLUEGRAY,
        "list_left": 0.187, "list_text": 0.375, "list_hanging": 0.188,
        "list_after": 4, "list_line": 1.25,
    },
    "compact_reference_guide": {
        "body_after": 6, "line": 1.25, "align": WD_ALIGN_PARAGRAPH.RIGHT,
        "h1_before": 18, "h1_after": 10, "h2_before": 14, "h2_after": 7,
        "h3_before": 10, "h3_after": 5, "table_fill": BLUEGRAY,
        "list_left": 0.187, "list_text": 0.375, "list_hanging": 0.188,
        "list_after": 4, "list_line": 1.25,
    },
    "launch_messaging_guide": {
        "body_after": 6, "line": 1.25, "align": WD_ALIGN_PARAGRAPH.RIGHT,
        "h1_before": 18, "h1_after": 10, "h2_before": 14, "h2_after": 7,
        "h3_before": 10, "h3_after": 5, "table_fill": BLUEGRAY,
        "list_left": 0.187, "list_text": 0.375, "list_hanging": 0.188,
        "list_after": 4, "list_line": 1.25,
    },
}


def rgb(hexstr):
    return RGBColor.from_string(hexstr)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D7DDD9", size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def set_table_geometry(table, widths_dxa, indent_dxa=120, total_dxa=9360):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    bidi = tbl_pr.find(qn("w:bidiVisual"))
    if bidi is None:
        bidi = OxmlElement("w:bidiVisual")
        tbl_pr.append(bidi)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            width = widths_dxa[min(i, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, color=INK, bold=False, italic=False, font="Arial"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic
    rtl = run._element.get_or_add_rPr().find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        run._element.get_or_add_rPr().append(rtl)
    return run


def set_para_rtl(p, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p.alignment = align
    p_pr = p._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    return p


def keep_with_next(p):
    p.paragraph_format.keep_with_next = True
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run("صفحة ")
    set_run(run, size=9, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def setup_styles(doc, preset_name):
    preset = PRESETS[preset_name]
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(preset["body_after"])
    normal.paragraph_format.line_spacing = preset["line"]
    normal.paragraph_format.alignment = preset["align"]

    for style_name, size, color, before, after in [
        ("Heading 1", 16, GREEN, preset["h1_before"], preset["h1_after"]),
        ("Heading 2", 13, GREEN2, preset["h2_before"], preset["h2_after"]),
        ("Heading 3", 12, INK, preset["h3_before"], preset["h3_after"]),
    ]:
        st = doc.styles[style_name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        st._element.rPr.rFonts.set(qn("w:cs"), "Arial")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = rgb(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    header = sec.header
    hp = header.paragraphs[0]
    hp.clear()
    set_para_rtl(hp, WD_ALIGN_PARAGRAPH.RIGHT)
    set_run(hp.add_run("نبري | نظام المبيعات والتوزيع"), size=9, color=GRAY, bold=True)
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.clear()
    add_page_number(fp)
    return preset


def add_title(doc, text_value, size=30, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER, after=8):
    p = doc.add_paragraph()
    set_para_rtl(p, align)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    set_run(p.add_run(text_value), size=size, color=color, bold=True)
    return p


def add_subtitle(doc, text_value, size=14, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12, bold=False):
    p = doc.add_paragraph()
    set_para_rtl(p, align)
    p.paragraph_format.space_after = Pt(after)
    set_run(p.add_run(text_value), size=size, color=color, bold=bold)
    return p


def add_kicker(doc, text_value, color=GOLD, after=4):
    p = doc.add_paragraph()
    set_para_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(after)
    set_run(p.add_run(text_value), size=10, color=color, bold=True)
    return p


def add_body(doc, text_value, bold=False, color=INK, after=None, align=WD_ALIGN_PARAGRAPH.RIGHT, italic=False):
    p = doc.add_paragraph()
    set_para_rtl(p, align)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    set_run(p.add_run(text_value), size=11, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text_value, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_para_rtl(p, WD_ALIGN_PARAGRAPH.RIGHT)
    set_run(p.add_run(text_value), size={1:16,2:13,3:12}[level], color={1:GREEN,2:GREEN2,3:INK}[level], bold=True)
    return p


def ensure_numbering(doc, preset):
    numbering = doc.part.numbering_part.element
    max_abs = max([int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))] or [0])
    max_num = max([int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))] or [0])
    ids = {}
    for kind, fmt, text_value in (("bullet", "bullet", "•"), ("decimal", "decimal", "%1.")):
        max_abs += 1
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(max_abs))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text_value)
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "right")
        lvl.append(lvl_jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(int(preset["list_text"] * 1440)))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:start"), str(int(preset["list_text"] * 1440)))
        ind.set(qn("w:hanging"), str(int(preset["list_hanging"] * 1440)))
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), str(int(preset["list_after"] * 20)))
        spacing.set(qn("w:line"), str(int(preset["list_line"] * 240)))
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Arial")
        r_fonts.set(qn("w:hAnsi"), "Arial")
        r_pr.append(r_fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        max_num += 1
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(max_num))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(max_abs))
        num.append(abs_id)
        numbering.append(num)
        ids[kind] = max_num
    return ids


def add_list(doc, items, num_id, numbered=False):
    for item in items:
        p = doc.add_paragraph()
        set_para_rtl(p, WD_ALIGN_PARAGRAPH.RIGHT)
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_el)
        p_pr.append(num_pr)
        set_run(p.add_run(item), size=11, color=INK)
    return p


def add_callout(doc, title, body, fill=LIME2, accent=GREEN):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_para_rtl(p, WD_ALIGN_PARAGRAPH.RIGHT)
    set_run(p.add_run(title + "\n"), size=12, color=accent, bold=True)
    set_run(p.add_run(body), size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_info_table(doc, rows, widths=(2700, 6660), header=None, header_fill=GREEN):
    count = len(rows) + (1 if header else 0)
    table = doc.add_table(rows=count, cols=2)
    set_table_geometry(table, list(widths), indent_dxa=120)
    set_table_borders(table)
    rindex = 0
    if header:
        table.cell(0, 0).merge(table.cell(0, 1))
        cell = table.cell(0, 0)
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        set_para_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
        set_run(p.add_run(header), size=11, color=WHITE, bold=True)
        rindex = 1
    for i, (label, value) in enumerate(rows, start=rindex):
        c0, c1 = table.rows[i].cells
        set_cell_shading(c0, "F2F4F7")
        p0, p1 = c0.paragraphs[0], c1.paragraphs[0]
        set_para_rtl(p0, WD_ALIGN_PARAGRAPH.RIGHT)
        set_para_rtl(p1, WD_ALIGN_PARAGRAPH.RIGHT)
        set_run(p0.add_run(label), size=10.5, color=GREEN, bold=True)
        set_run(p1.add_run(value), size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_matrix(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_geometry(table, widths_dxa, indent_dxa=120)
    set_table_borders(table)
    for i, head in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, GREEN)
        p = cell.paragraphs[0]
        set_para_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
        set_run(p.add_run(head), size=10, color=WHITE, bold=True)
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            cell = table.rows[r].cells[c]
            if r % 2 == 0:
                set_cell_shading(cell, "F7F9F8")
            p = cell.paragraphs[0]
            set_para_rtl(p, WD_ALIGN_PARAGRAPH.RIGHT if c else WD_ALIGN_PARAGRAPH.CENTER)
            set_run(p.add_run(str(value)), size=9.5, color=INK, bold=(c == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_screenshot(doc, filename, caption):
    path = ASSETS / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.2))
    cap = doc.add_paragraph()
    set_para_rtl(cap, WD_ALIGN_PARAGRAPH.CENTER)
    cap.paragraph_format.space_after = Pt(12)
    set_run(cap.add_run(caption), size=9, color=GRAY, italic=True)


def add_cover(doc, kicker, title, subtitle, metadata, pattern="proposal", screenshot=None):
    for _ in range(4 if pattern != "editorial" else 7):
        doc.add_paragraph()
    add_kicker(doc, kicker, color=GOLD, after=12)
    add_title(doc, title, size=30 if pattern != "customer" else 28, color=GREEN, after=8)
    add_subtitle(doc, subtitle, size=14, color=GRAY, after=22)
    if metadata:
        table = doc.add_table(rows=len(metadata), cols=2)
        set_table_geometry(table, [2800, 6560], indent_dxa=120)
        remove_table_borders(table)
        for i, (label, value) in enumerate(metadata):
            c0, c1 = table.rows[i].cells
            set_cell_shading(c0, LIME2)
            set_cell_shading(c1, "F7F9F8")
            for cell in (c0, c1):
                set_cell_margins(cell, top=90, start=140, bottom=90, end=140)
            p0, p1 = c0.paragraphs[0], c1.paragraphs[0]
            set_para_rtl(p0, WD_ALIGN_PARAGRAPH.RIGHT)
            set_para_rtl(p1, WD_ALIGN_PARAGRAPH.RIGHT)
            set_run(p0.add_run(label), size=10, color=GREEN, bold=True)
            set_run(p1.add_run(value), size=10, color=INK)
    if screenshot:
        doc.add_paragraph()
        add_screenshot(doc, screenshot, "لقطة من النسخة التشغيلية للمرحلة الأولى")
    doc.add_page_break()


def add_signature_block(doc):
    table = doc.add_table(rows=6, cols=2)
    set_table_geometry(table, [4680, 4680], indent_dxa=120)
    set_table_borders(table, color="AEB8B2", size=6)
    titles = ["الطرف الأول - مقدم النظام", "الطرف الثاني - العميل"]
    for c, title in enumerate(titles):
        set_cell_shading(table.cell(0, c), GREEN)
        p = table.cell(0, c).paragraphs[0]
        set_para_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
        set_run(p.add_run(title), size=10.5, color=WHITE, bold=True)
    labels = ["الاسم:", "الصفة:", "التوقيع:", "التاريخ:", "الختم:"]
    for r, label in enumerate(labels, start=1):
        for c in range(2):
            p = table.cell(r, c).paragraphs[0]
            set_para_rtl(p, WD_ALIGN_PARAGRAPH.RIGHT)
            set_run(p.add_run(label + " ______________________________"), size=10, color=INK)
    return table


def start_doc(preset_name):
    doc = Document()
    preset = setup_styles(doc, preset_name)
    nums = ensure_numbering(doc, preset)
    return doc, preset, nums


def build_proposal():
    doc, preset, nums = start_doc("narrative_proposal")
    add_cover(
        doc,
        "عرض تجاري وتقني",
        "نظام نبري للمبيعات والتوزيع",
        "من الطلب إلى التحصيل في مسار واحد قابل للتوسع",
        [
            ("مقدم إلى", "شركة نبري - السيد غسان جيلاني"),
            ("نسخة العرض", "المرحلة الأولى - V1"),
            ("التاريخ", "21 أغسطس 2026"),
            ("حالة الوثيقة", "مسودة تجارية قابلة للتخصيص"),
        ],
        pattern="proposal",
        screenshot="01-admin-dashboard.png",
    )
    add_callout(doc, "الخلاصة التنفيذية", "نظام نبري هو قاعدة تشغيل للمبيعات تربط الإدارة والمندوب والعميل بقائمة أسعار واحدة وطلبات وتحصيلات قابلة للمتابعة. المرحلة الأولى جاهزة للتجربة المنضبطة قبل أي توسع أو أتمتة إضافية.")
    add_heading(doc, "1. المشكلة التي يعالجها النظام", 1)
    add_body(doc, "عندما تكون الأسعار والطلبات والتحصيلات موزعة بين ملفات ورسائل واتصالات، يصبح الخطأ في السعر أو حالة الطلب أو رصيد العميل احتمالًا يوميًا. النظام يجمع هذه الخطوات في دورة بيع واحدة دون محاولة تحويل الشركة إلى منظومة ضخمة من اليوم الأول.")
    add_list(doc, [
        "توحيد السعر المعروض للإدارة والمندوب والعميل.",
        "تسجيل الطلب ومتابعة حالته من المراجعة إلى التجهيز والتسليم.",
        "ربط التحصيل بتاريخ العميل ورصيده.",
        "إظهار أولويات الزيارة وفرص إعادة الطلب للمندوب.",
        "حماية المرحلة الأولى من التوسع غير الضروري عبر تأجيل الوظائف غير الجاهزة بمتطلبات واضحة.",
    ], nums["bullet"])

    add_heading(doc, "2. ما تم إنجازه في المرحلة الأولى", 1)
    add_matrix(doc, ["الواجهة", "الوظائف المتاحة", "القيمة التشغيلية"], [
        ("الإدارة", "مركز القيادة، العملاء، الطلبات، المسارات، التحصيل، الأسعار والمخزون، الهيكل", "رقابة وقرار من شاشة واحدة"),
        ("المندوب / الموظف", "خطة الزيارات، أولويات العملاء، الطلبات والتحصيلات، تحديث الحالة", "تنفيذ يومي منظم وقابل للمتابعة"),
        ("العميل", "كتالوج، أسعار موحدة، سلة، إرسال طلب، متابعة الطلبات", "طلب أسهل وأقل اعتمادًا على النقل اليدوي"),
    ], [1600, 4560, 3200])
    add_screenshot(doc, "02-price-list.png", "قائمة الأسعار والمخزون داخل واجهة الإدارة")

    add_heading(doc, "3. مصدر السعر الواحد", 1)
    add_body(doc, "تعتمد المرحلة الأولى على قائمة أسعار موحدة برقم إصدار NAPRI-P1-2026-08. السعر الظاهر في واجهة العميل هو نفسه السعر المتاح للمندوب والإدارة، ويعيد النظام التحقق من السعر والمخزون قبل حفظ الطلب.")
    add_info_table(doc, [
        ("نسخة القائمة", "NAPRI-P1-2026-08"),
        ("حالة الاعتماد", "مسودة اعتماد الإدارة"),
        ("نقطة التحكم", "قائمة الأسعار والمخزون في واجهة الإدارة"),
        ("التحقق", "الخادم يعيد التحقق من السعر والكمية قبل التسجيل"),
    ], header="ضبط الأسعار")

    add_heading(doc, "4. دورة البيع المستهدفة", 1)
    add_list(doc, [
        "اختيار المنتجات بواسطة العميل أو المندوب.",
        "تسجيل الطلب بعد التحقق من السعر والمخزون.",
        "مراجعة الإدارة وتحديث حالة الطلب.",
        "تجهيز الطلب وتسليمه ومتابعة التنفيذ.",
        "تسجيل التحصيل وتحديث رصيد العميل.",
    ], nums["decimal"], numbered=True)
    add_screenshot(doc, "03-orders.png", "شاشة الطلبات والمبيعات بعد ربط التسجيل بحالة التنفيذ")

    add_heading(doc, "5. نموذج التشغيل التجريبي المقترح", 1)
    add_matrix(doc, ["الفترة", "العمل", "مؤشر القبول"], [
        ("اليوم 1", "اعتماد قائمة الأسعار ومسؤوليات التشغيل", "قائمة معتمدة ومالكا عملية واضحان"),
        ("الأيام 2-3", "تجهيز أول 20 عميلًا وتدريب المسؤولين", "بيانات أساسية صالحة للاستخدام"),
        ("الأيام 4-8", "تشغيل طلبات فعلية محدودة ومتابعة الأخطاء", "طلبات وحالات وتحصيلات مسجلة"),
        ("اليومان 9-10", "مراجعة النتائج وتثبيت التحسينات", "قرار واضح بالاستمرار أو التعديل"),
    ], [1500, 4860, 3000])

    add_heading(doc, "6. نموذج تجاري قابل للتخصيص", 1)
    add_callout(doc, "لا توجد أرقام نهائية في هذه النسخة", "تُستكمل الأسعار والضرائب ومواعيد الدفع بعد تحديد الجهة البائعة، الدولة، نطاق الاستضافة، مستوى الدعم، وعدد المستخدمين.", fill="FFF6E5", accent=GOLD)
    add_info_table(doc, [
        ("رسوم الإعداد والتنفيذ", "[يُحدد المبلغ والعملة والضرائب]"),
        ("الاشتراك أو الصيانة", "[شهري / سنوي / بدون اشتراك - يُحدد]"),
        ("الاستضافة", "[مشمولة / على حساب العميل / بند منفصل]"),
        ("الدعم", "[الساعات والقنوات وزمن الاستجابة]"),
        ("التعديلات خارج النطاق", "تُنفذ بطلب تغيير وعرض سعر مستقل"),
    ], header="البنود المالية")

    add_heading(doc, "7. ما ليس ضمن المرحلة الأولى", 1)
    add_list(doc, [
        "التسويق الآلي والحملات؛ يحتاج قنوات معتمدة ومحتوى وسياسات إرسال.",
        "التتبع الجغرافي المباشر؛ يحتاج عناوين وإحداثيات وموافقة تشغيلية.",
        "نظام صلاحيات متقدم؛ يحتاج تعريف المستخدمين والأدوار وسياسات الوصول.",
        "محاسبة وضرائب كاملة؛ تحتاج متطلبات محاسبية وقانونية منفصلة.",
        "تكامل واتساب API أو بوابات دفع؛ يخضع لحسابات وموافقات ورسوم أطراف خارجية.",
    ], nums["bullet"])

    add_heading(doc, "8. معايير القبول", 1)
    add_list(doc, [
        "تعمل الواجهات الثلاث وتفتح الوظائف الداخلة في النطاق.",
        "يظهر السعر نفسه للمنتج عبر الواجهات ويُراجع عند حفظ الطلب.",
        "يظهر طلب العميل في واجهة الإدارة ويمكن تحديث حالته.",
        "يمكن تسجيل تحصيل وتحديث رصيد العميل.",
        "توضح الوظائف المؤجلة متطلبات تفعيلها بدل إظهار عناصر غير مستجيبة.",
    ], nums["bullet"])

    add_heading(doc, "9. القرار المطلوب", 1)
    add_callout(doc, "بدء تشغيل محدود", "اعتماد قائمة الأسعار، اختيار أول 20 عميلًا، وتحديد مسؤول مراجعة الطلبات والتحصيلات يوميًا. بعد عشرة أيام تشغيل تُراجع النتائج قبل فتح نطاق جديد.")
    add_body(doc, "صلاحية العرض التجاري: [عدد الأيام من تاريخ الإصدار]. لا يُعد هذا العرض عقدًا ملزمًا ما لم يُوقع اتفاق مستقل يحدد الأطراف والسعر والنطاق والمسؤوليات.", italic=True, color=GRAY)
    path = OUT / "01-العرض-التجاري-لنظام-نبري.docx"
    doc.save(path)
    return path


def build_contract():
    doc, preset, nums = start_doc("contract_negotiation_brief")
    add_cover(
        doc,
        "مسودة اتفاقية - غير جاهزة للتوقيع قبل الاستكمال القانوني",
        "اتفاقية توريد وترخيص وتشغيل نظام نبري",
        "مسودة عامة لتنظيم النطاق والترخيص والاستضافة والدعم والبيانات",
        [
            ("الطرف الأول", "[الاسم القانوني لمقدم النظام]"),
            ("الطرف الثاني", "[الاسم القانوني للعميل / شركة نبري]"),
            ("الدولة والقانون", "[يُحدد قبل التوقيع]"),
            ("الإصدار", "DRAFT 0.1 - 21 أغسطس 2026"),
        ],
        pattern="memo",
    )
    add_callout(doc, "تنبيه قانوني", "هذه وثيقة إعداد تجاري وتقني وليست استشارة قانونية. يجب مراجعتها بواسطة محامٍ مرخص في الدولة المختصة قبل التوقيع، مع استكمال جميع الحقول بين الأقواس.", fill="FFF1F0", accent=RED)

    add_heading(doc, "بيانات الأطراف", 1)
    add_info_table(doc, [
        ("الطرف الأول", "[الاسم القانوني، الشكل القانوني، رقم التسجيل، العنوان، الممثل وصفته]"),
        ("الطرف الثاني", "[الاسم القانوني، الشكل القانوني، رقم التسجيل، العنوان، الممثل وصفته]"),
        ("تاريخ النفاذ", "[اليوم / الشهر / السنة]"),
        ("البريد المعتمد للإشعارات", "[بريد الطرف الأول] / [بريد الطرف الثاني]"),
    ], header="معلومات التعاقد")

    clauses = [
        ("1. التمهيد والتعاريف", [
            "يُعد التمهيد والملاحق جزءًا من الاتفاقية. ويقصد بالنظام: تطبيق نبري للمبيعات والتوزيع بواجهات الإدارة والمندوب أو الموظف والعميل، وفق النطاق المبين في الملحق (أ).",
            "يقصد ببيانات العميل كل البيانات التجارية أو الشخصية أو التشغيلية التي يدخلها الطرف الثاني أو مستخدموه. ويقصد بطلب التغيير أي عمل خارج النطاق المتفق عليه.",
        ]),
        ("2. موضوع الاتفاقية ونطاقها", [
            "يلتزم الطرف الأول بتوفير وترخيص وتشغيل المرحلة الأولى من النظام للطرف الثاني، وتنفيذ أعمال التهيئة والتدريب والدعم المحددة في الاتفاقية وملاحقها.",
            "أي ميزة غير مدرجة صراحة في الملحق (أ) تُعد خارج النطاق ولا تصبح التزامًا إلا بطلب تغيير مكتوب يحدد السعر والمدة ومعيار القبول.",
        ]),
        ("3. الترخيص والاستخدام", [
            "يمنح الطرف الأول للطرف الثاني ترخيصًا [غير حصري، غير قابل للتنازل] لاستخدام النظام داخل أعماله طوال مدة الاتفاقية ولعدد [يُحدد] من المستخدمين.",
            "لا يجوز نسخ الشفرة المصدرية أو إعادة بيع النظام أو منح حق استخدامه لطرف ثالث أو إزالة إشعارات الملكية إلا بموافقة مكتوبة.",
            "تظل البيانات التي يدخلها الطرف الثاني ملكًا له، بينما تظل الشفرة والتصميمات والمكونات العامة وحقوق التطوير ملكًا للطرف الأول ما لم يتفق كتابة على خلاف ذلك.",
        ]),
        ("4. التنفيذ والتسليم والقبول", [
            "يبدأ التنفيذ بعد استلام الدفعة الأولية والبيانات المطلوبة واعتماد قائمة الأسعار ومسؤولي التشغيل.",
            "تكون فترة القبول [عدد] أيام عمل من تاريخ الإتاحة. ويقدم الطرف الثاني خلالها ملاحظاته المرتبطة بمعايير القبول في الملحق (ب).",
            "تُعالج العيوب التي تمنع وظيفة متفقًا عليها دون رسوم إضافية. أما التحسينات أو التغييرات في إجراءات العمل فتخضع لطلب تغيير.",
        ]),
        ("5. المقابل المالي والدفع", [
            "رسوم الإعداد والتنفيذ: [المبلغ والعملة]، والضرائب: [مشمولة / غير مشمولة].",
            "رسوم الاستضافة أو الاشتراك أو الصيانة: [المبلغ والدورية].",
            "جدول الدفع: [نسبة عند التوقيع]، [نسبة عند الإتاحة]، [نسبة عند القبول].",
            "في حال التأخر عن السداد أكثر من [عدد] يومًا بعد إخطار كتابي، يجوز تعليق الخدمات غير الحرجة مع منح فرصة معقولة لتصدير البيانات، وذلك وفق القانون المختص.",
        ]),
        ("6. الاستضافة والتوافر والنسخ الاحتياطي", [
            "تكون الاستضافة [بواسطة الطرف الأول / حساب الطرف الثاني / مزود خارجي]، وتحدد بيئة الإنتاج والدولة أو المنطقة المستضيفة قبل بدء التشغيل.",
            "مستهدف التوافر الشهري: [نسبة] باستثناء الصيانة المعلنة والقوة القاهرة وأعطال مزودي الخدمات الخارجيين.",
            "سياسة النسخ الاحتياطي والاستعادة: [التكرار، مدة الاحتفاظ، زمن الاستعادة]، وتُختبر وفق جدول متفق عليه.",
        ]),
        ("7. الدعم والصيانة", [
            "قنوات الدعم: [بريد / واتساب أعمال / بوابة دعم]. ساعات الدعم: [يُحدد].",
            "الأولوية الحرجة: توقف شامل يمنع البيع؛ زمن الاستجابة المستهدف [يُحدد]. الأولوية العادية: عيب لا يمنع العمل؛ زمن الاستجابة [يُحدد].",
            "لا تشمل الصيانة تغييرات الأنظمة الخارجية أو الأسعار أو سياسات مزودي الخدمات إلا وفق طلب تغيير.",
        ]),
        ("8. حماية البيانات والسرية", [
            "يلتزم كل طرف بحماية المعلومات السرية وعدم استخدامها إلا لتنفيذ الاتفاقية، مع قصر الوصول على من يحتاج إليها.",
            "يحدد الطرفان أدوارهما القانونية بشأن البيانات الشخصية، وأساس المعالجة، ومدة الاحتفاظ، وإجراءات الحوادث، وطلبات أصحاب البيانات وفق قانون الدولة المختصة.",
            "يخطر الطرف المتأثر الآخر بحادث أمني جوهري دون تأخير غير مبرر ووفق المواعيد التي يفرضها القانون المختص.",
        ]),
        ("9. الملكية الفكرية", [
            "يحتفظ كل طرف بحقوقه السابقة على الاتفاقية. ولا تنتقل ملكية النظام أو مكوناته إلا بنص صريح ومقابل محدد.",
            "يجوز للطرف الأول استخدام المعرفة والخبرة العامة غير المتضمنة لبيانات العميل أو أسراره في مشروعات أخرى.",
            "تُحدد مكونات الأطراف الثالثة وتراخيصها وشروط استخدامها عند وجودها.",
        ]),
        ("10. الضمانات وحدود المسؤولية", [
            "يضمن الطرف الأول أن النظام سيؤدي الوظائف المتفق عليها جوهريًا عند استخدامه وفق الدليل، مع معالجة العيوب المثبتة خلال فترة الضمان [المدة].",
            "لا يضمن النظام نتيجة تجارية محددة، ولا يتحمل الطرف الأول أخطاء البيانات المدخلة أو الأسعار غير المعتمدة أو انقطاع خدمات الغير.",
            "يكون الحد الأقصى للمسؤولية المباشرة [المبلغ أو نسبة الرسوم]، مع مراعاة الاستثناءات التي لا يسمح القانون المختص بتحديدها أو استبعادها.",
        ]),
        ("11. المدة والإنهاء والخروج", [
            "مدة الاتفاقية [المدة] وتجدد [تلقائيًا / بموافقة مكتوبة].",
            "يجوز الإنهاء عند إخلال جوهري لا يُعالج خلال [عدد] يومًا من الإخطار، أو عند الإعسار، أو وفق أي حق إلزامي في القانون المختص.",
            "عند الانتهاء، يتيح الطرف الأول تصدير بيانات الطرف الثاني بصيغة [CSV / Excel / أخرى] خلال [المدة]، ثم يحذف النسخ المتبقية وفق سياسة الاحتفاظ والقانون.",
        ]),
        ("12. الإشعارات والتوقيع الإلكتروني", [
            "ترسل الإشعارات إلى العناوين والبريد المعتمدين في صدر الاتفاقية، ويحتفظ الطرفان بسجل الإرسال والاستلام.",
            "يجوز توقيع الاتفاقية إلكترونيًا أو في نسخ متقابلة إذا كان ذلك معترفًا به في القانون المختص واستوفى متطلبات تحديد هوية الموقع وسلامة السجل.",
        ]),
        ("13. القانون وتسوية النزاعات", [
            "تخضع الاتفاقية لقوانين [الدولة / الولاية].",
            "يسعى الطرفان للتسوية الودية خلال [عدد] يومًا، ثم يكون الاختصاص لـ[المحاكم المختصة / التحكيم ومكانه وقواعده ولغته].",
        ]),
        ("14. أحكام عامة", [
            "لا يجوز التنازل عن الاتفاقية دون موافقة الطرف الآخر، باستثناء إعادة الهيكلة أو نقل النشاط وفق ما يسمح به القانون.",
            "تعد الاتفاقية وملاحقها كامل التفاهم، ولا يكون التعديل نافذًا إلا مكتوبًا وموقعًا من ممثلين مفوضين.",
            "إذا تعذر تنفيذ بند، يبقى باقي الاتفاق صحيحًا ويستبدل البند بما يحقق غرضه المشروع قدر الإمكان.",
        ]),
    ]
    for title, paragraphs in clauses:
        add_heading(doc, title, 1)
        for paragraph in paragraphs:
            add_body(doc, paragraph)

    add_heading(doc, "الملحق (أ) - نطاق المرحلة الأولى", 1)
    add_matrix(doc, ["الوحدة", "داخل النطاق", "خارج النطاق أو مشروط"], [
        ("الإدارة", "العملاء، الطلبات، الحالات، التحصيل، الأسعار، المخزون الأساسي، المسارات", "محاسبة كاملة وصلاحيات مؤسسية متقدمة"),
        ("المندوب", "زيارات، أولويات، طلبات وتحصيلات وتحديث حالة", "GPS مباشر دون عناوين وإحداثيات معتمدة"),
        ("العميل", "كتالوج وسلة وطلب ومتابعة", "دفع إلكتروني أو حسابات دخول دون اتفاق وتكامل"),
        ("الأسعار", "قائمة موحدة وتحقق عند الحفظ", "سياسات خصم وائتمان متعددة المستويات"),
        ("التكاملات", "رابط واتساب مبسط عند اعتماده", "WhatsApp API وبوابات الدفع وخدمات الغير"),
    ], [1550, 4410, 3400])

    add_heading(doc, "الملحق (ب) - اختبار القبول", 1)
    add_list(doc, [
        "فتح الواجهات الثلاث والوظائف المتفق عليها دون خطأ مانع.",
        "تطابق سعر المنتج في الإدارة والمندوب والعميل.",
        "إنشاء طلب وحفظه وظهوره في الإدارة.",
        "تحديث حالة الطلب وحفظ التغيير.",
        "تسجيل تحصيل وتحديث رصيد العميل.",
        "إظهار متطلبات الوظائف المؤجلة عند فتحها.",
    ], nums["bullet"])

    add_heading(doc, "التوقيعات", 1)
    add_signature_block(doc)
    add_heading(doc, "ملاحظة إعداد خارج متن الاتفاقية", 1)
    add_body(doc, "رُوعي في إعداد هذه المسودة مبدأ وضوح نطاق الترخيص، ملكية البيانات، القبول، الدعم، الخروج، والتوقيع الإلكتروني. يجب مواءمة الصياغة مع القانون الوطني. من المراجع العامة: قانون الأونسيترال النموذجي للتجارة الإلكترونية، ودليل الويبو لترخيص التكنولوجيا.", color=GRAY, italic=True)
    add_body(doc, "المراجع: https://uncitral.un.org/en/texts/ecommerce/modellaw/electronic_commerce | https://www.wipo.int/publications/en/details.jsp?id=296&plang=EN", color=GRAY, align=WD_ALIGN_PARAGRAPH.LEFT)
    path = OUT / "02-مسودة-اتفاقية-توريد-وترخيص-نظام-نبري.docx"
    doc.save(path)
    return path


def build_manual():
    doc, preset, nums = start_doc("compact_reference_guide")
    add_cover(
        doc,
        "دليل التشغيل والاستخدام",
        "نظام نبري - المرحلة الأولى",
        "دليل الإدارة والمندوب والعميل من فتح النظام إلى تسجيل التحصيل",
        [
            ("إصدار النظام", "V1 - المرحلة الأولى"),
            ("قائمة الأسعار", "NAPRI-P1-2026-08"),
            ("الجمهور", "الإدارة، المندوبون والموظفون، العملاء"),
            ("تاريخ الدليل", "21 أغسطس 2026"),
        ],
        pattern="editorial",
    )
    add_callout(doc, "فكرة الدليل", "استخدم القسم الخاص بدورك يوميًا، وارجع إلى مسار البيع الكامل عند اختبار الربط بين الواجهات.")
    add_heading(doc, "1. قبل بدء الاستخدام", 1)
    add_list(doc, [
        "افتح النظام من جهاز متصل بالإنترنت أو من النسخة المحلية أثناء الاختبار.",
        "تأكد أن مؤشر قاعدة البيانات يظهر متصلة.",
        "تحقق من رقم إصدار قائمة الأسعار قبل إنشاء أي طلب.",
        "لا تستخدم بيانات عملاء حقيقية حساسة في نسخة العرض المؤقتة العامة إن تم نشرها لاحقًا.",
    ], nums["bullet"])
    add_info_table(doc, [
        ("واجهة الإدارة", "الصورة الكاملة والقرارات والمتابعة"),
        ("واجهة المندوب والموظف", "التنفيذ اليومي والزيارات والطلبات والتحصيل"),
        ("واجهة العميل", "الكتالوج والسلة وإرسال الطلب والمتابعة"),
    ], header="اختيار الواجهة")
    add_screenshot(doc, "01-admin-dashboard.png", "مركز قيادة الإدارة والانتقال بين الواجهات")

    add_heading(doc, "2. دليل واجهة الإدارة", 1)
    add_heading(doc, "2.1 مركز القيادة", 2)
    add_body(doc, "تعرض الشاشة مؤشرات المبيعات الحالية، الرصيد المطلوب تحصيله، فرص إعادة الطلب، المنتجات تحت حد الأمان، والمسارات المقترحة. استخدمها لبدء اليوم وتحديد الأولويات.")
    add_heading(doc, "2.2 العملاء والتوقعات", 2)
    add_list(doc, [
        "راجع بيانات العميل ورصيده والمندوب المسؤول.",
        "استخدم توقع إعادة الطلب لترتيب الاتصال أو الزيارة.",
        "لا تعتبر التوقع التزامًا؛ هو أداة ترتيب تحتاج مراجعة بشرية.",
    ], nums["bullet"])
    add_heading(doc, "2.3 الطلبات والمبيعات", 2)
    add_list(doc, [
        "افتح قائمة الطلبات وشاهد المصدر والعميل والإجمالي.",
        "غيّر الحالة إلى مراجعة أو قيد التجهيز أو تم التسليم حسب الواقع.",
        "لا تسجل التسليم قبل التأكد من تنفيذ العملية فعليًا.",
    ], nums["decimal"])
    add_screenshot(doc, "03-orders.png", "شاشة الطلبات وتحديث حالة التنفيذ")
    add_heading(doc, "2.4 قائمة الأسعار والمخزون", 2)
    add_body(doc, "هذه الشاشة هي مصدر السعر الواحد. راجع المنتج والسعر والكمية المتاحة وحد الأمان. أي تعديل للسعر يجب أن يصدر كنسخة معتمدة وواضحة بدل تغيير غير موثق.")
    add_screenshot(doc, "02-price-list.png", "قائمة الأسعار والمخزون وربطها بواجهة العميل")
    add_heading(doc, "2.5 التحصيل والديون", 2)
    add_list(doc, [
        "اختر العميل الصحيح.",
        "أدخل مبلغ التحصيل ووسيلته والمرجع إن وجد.",
        "راجع الرصيد الجديد بعد الحفظ.",
        "احتفظ بالمستند المالي الخارجي عند الحاجة؛ النظام في المرحلة الأولى سجل متابعة وليس نظام محاسبة قانونيًا كاملًا.",
    ], nums["decimal"])
    add_heading(doc, "2.6 خطوط السير", 2)
    add_body(doc, "تظهر المسارات والأولويات التشغيلية. وظائف الخرائط أو GPS المؤجلة تعرض متطلبات التفعيل مثل العناوين والإحداثيات والموافقات التشغيلية.")
    add_screenshot(doc, "04-routes.png", "خطوط السير وأولويات الزيارة")

    add_heading(doc, "3. دليل المندوب أو الموظف", 1)
    add_screenshot(doc, "05-employee.png", "واجهة التنفيذ اليومي للمندوب أو الموظف")
    add_list(doc, [
        "ابدأ بالعملاء ذوي الأولوية أو التحصيل المستحق.",
        "افتح العميل وراجع آخر حركة ورصيد الطلب المتوقع.",
        "سجل الطلب من قائمة الأسعار الحالية ولا تعتمد سعرًا من رسالة قديمة.",
        "حدّث تقدم الزيارة والحالة بعد التنفيذ الفعلي.",
        "سجل التحصيل بدقة ثم راجع الرصيد.",
    ], nums["decimal"])
    add_callout(doc, "قاعدة المندوب", "كل سعر من القائمة الحالية، وكل حالة من الواقع، وكل تحصيل بمبلغ ومرجع يمكن مراجعته.")

    add_heading(doc, "4. دليل واجهة العميل", 1)
    add_screenshot(doc, "06-customer.png", "متجر نبري والمنتجات المرتبطة بقائمة الأسعار الموحدة")
    add_list(doc, [
        "تأكد من رقم نسخة قائمة الأسعار الظاهر أعلى المنتجات.",
        "استخدم زر الزيادة لاختيار الكمية المطلوبة.",
        "راجع عدد الوحدات والإجمالي في سلة الطلب.",
        "أرسل الطلب للتأكيد؛ يعيد النظام فحص السعر والمخزون قبل التسجيل.",
        "تابع الطلب من قسم طلباتي.",
    ], nums["decimal"])
    add_screenshot(doc, "07-customer-cart.png", "سلة العميل بعد اختيار منتج من القائمة الموحدة")

    add_heading(doc, "5. اختبار الربط الكامل", 1)
    add_matrix(doc, ["الخطوة", "المنفذ", "النتيجة المتوقعة"], [
        ("1", "العميل يختار منتجًا ويرسل الطلب", "يُحفظ الطلب بالسعر والكمية المتاحة"),
        ("2", "الإدارة تفتح الطلبات", "يظهر الطلب الجديد ومصدره وإجماليه"),
        ("3", "الإدارة تحدث الحالة", "تُحفظ الحالة وتظهر في المتابعة"),
        ("4", "المندوب ينفذ أو يتابع", "يُحدَّث تقدم العمل"),
        ("5", "المحاسب أو الإدارة يسجل التحصيل", "ينخفض رصيد العميل"),
    ], [1000, 4000, 4360])

    add_heading(doc, "6. الوظائف المؤجلة ومتطلبات تفعيلها", 1)
    add_matrix(doc, ["الوظيفة", "سبب التأجيل", "المطلوب للتفعيل"], [
        ("التسويق الآلي", "ليست ضرورية للبيع الأول", "قنوات، محتوى، موافقات، سياسة إرسال"),
        ("GPS", "بيانات الموقع غير مكتملة", "عناوين وإحداثيات وموافقة تشغيلية"),
        ("الصلاحيات المتقدمة", "المستخدمون والأدوار لم تُعتمد", "قائمة مستخدمين ومصفوفة صلاحيات"),
        ("المخزون المتقدم", "إجراءات الجرد غير مثبتة", "جرد افتتاحي وسياسة حركة واعتماد"),
        ("واتساب API", "يتطلب حسابًا ومزودًا وموافقة", "حساب أعمال معتمد وقوالب وسياسة خصوصية"),
    ], [1900, 3000, 4460])

    add_heading(doc, "7. معالجة المشكلات الشائعة", 1)
    add_info_table(doc, [
        ("السعر غير متوقع", "راجع رقم نسخة القائمة ثم تواصل مع الإدارة قبل الإرسال"),
        ("زر الإرسال غير نشط", "تأكد من إضافة كمية ومن توفر مخزون"),
        ("الطلب لا يظهر", "تحقق من اتصال قاعدة البيانات، ثم أعد فتح الطلبات دون تكرار الإرسال مباشرة"),
        ("التحصيل لم يُحفظ", "لا تعتبر العملية مكتملة؛ راجع رسالة الخطأ والاتصال"),
        ("أيقونة تشرح متطلبات", "الوظيفة مؤجلة وليست معطلة؛ نفذ المتطلبات قبل طلب التفعيل"),
    ], header="تشخيص سريع")

    add_heading(doc, "8. قوائم الفحص اليومية", 1)
    add_heading(doc, "بداية اليوم - الإدارة", 2)
    add_list(doc, ["راجع اتصال قاعدة البيانات.", "راجع الطلبات الجديدة.", "راجع التحصيلات المستحقة.", "راجع المنتجات تحت حد الأمان.", "اعتمد أولويات المسارات."], nums["bullet"])
    add_heading(doc, "نهاية اليوم - الإدارة والمندوب", 2)
    add_list(doc, ["تحديث حالات الطلبات.", "تسجيل التحصيلات المنفذة.", "مراجعة الرصيد والاستثناءات.", "توثيق أي سعر أو مخزون يحتاج تعديلًا.", "تجميع الملاحظات للتحسين الأسبوعي."], nums["bullet"])

    add_heading(doc, "9. حدود الدليل والدعم", 1)
    add_body(doc, "هذا الدليل يشرح المرحلة الأولى كما هي في 21 أغسطس 2026. يجب تحديثه عند اعتماد صلاحيات جديدة أو تكاملات أو تغيير جوهري في دورة البيع.")
    add_info_table(doc, [
        ("مسؤول تشغيل النظام", "[يُحدد]"),
        ("دعم الأعمال", "[يُحدد]"),
        ("الدعم التقني", "[يُحدد]"),
        ("قناة الإبلاغ", "[بريد / واتساب / نموذج دعم]"),
    ], header="بيانات الدعم")
    path = OUT / "03-دليل-استخدام-نظام-نبري.docx"
    doc.save(path)
    return path


def build_brochure():
    doc, preset, nums = start_doc("launch_messaging_guide")
    add_cover(
        doc,
        "كتيب تعريفي",
        "نبري: المبيعات تبدأ من بيانات واحدة",
        "إدارة ومندوب وعميل - في دورة بيع موحدة من الطلب إلى التحصيل",
        [
            ("المنتج", "نظام نبري للمبيعات والتوزيع"),
            ("الإصدار", "المرحلة الأولى V1"),
            ("مناسب لـ", "الشركات الصغيرة والمتوسطة وفرق التوزيع"),
            ("حالة النسخة", "تشغيل تجريبي منضبط"),
        ],
        pattern="customer",
        screenshot="01-admin-dashboard.png",
    )
    add_heading(doc, "ماذا يقدم نظام نبري؟", 1)
    add_body(doc, "يحوّل نظام نبري الطلبات والأسعار والتحصيلات ومسارات المندوبين من معلومات متفرقة إلى عملية واحدة يمكن رؤيتها ومتابعتها. يبدأ بالضروري للبيع، ويؤجل ما لا يلزم حتى تتوفر بياناته ومتطلباته.")
    add_callout(doc, "وعد المرحلة الأولى", "سعر واحد، طلب مسجل، حالة واضحة، تحصيل مرتبط بالعميل.")

    add_heading(doc, "ثلاث واجهات لنفس العمل", 1)
    add_matrix(doc, ["الإدارة", "المندوب أو الموظف", "العميل"], [
        ("قرارات ومؤشرات", "زيارات وأولويات", "كتالوج وسلة"),
        ("طلبات وحالات", "طلبات وتحصيلات", "إرسال ومتابعة"),
        ("أسعار ومخزون", "السعر الموحد", "السعر الموحد"),
    ], [3120, 3120, 3120])
    add_screenshot(doc, "06-customer.png", "واجهة العميل تعرض الأسعار والمنتجات من المصدر الموحد")

    add_heading(doc, "لماذا قائمة الأسعار مهمة؟", 1)
    add_body(doc, "تُستخدم القائمة نفسها في الإدارة والمندوب ومتجر العميل. وعند إرسال الطلب يراجع النظام السعر والمخزون قبل الحفظ، ما يقلل الاعتماد على رسائل أو ملفات قديمة.")
    add_info_table(doc, [
        ("الإصدار الحالي", "NAPRI-P1-2026-08"),
        ("الحالة", "مسودة اعتماد الإدارة"),
        ("مبدأ التشغيل", "المصدر الواحد للسعر"),
    ], header="قائمة الأسعار الموحدة")
    add_screenshot(doc, "02-price-list.png", "الإدارة تراجع السعر والمخزون من شاشة واحدة")

    add_heading(doc, "رحلة البيع", 1)
    add_list(doc, [
        "العميل أو المندوب يختار المنتجات.",
        "النظام يسجل الطلب بعد التحقق.",
        "الإدارة تراجع وتحدث الحالة.",
        "المندوب يتابع التجهيز أو التسليم.",
        "التحصيل يُسجل على رصيد العميل.",
    ], nums["decimal"])

    add_heading(doc, "ما يعمل الآن وما ينتظر؟", 1)
    add_matrix(doc, ["يعمل الآن", "يُفعل لاحقًا عند اكتمال متطلباته"], [
        ("الطلبات والأسعار والمخزون الأساسي", "التسويق الآلي بعد اعتماد القنوات والمحتوى"),
        ("العملاء والمسارات والتوقعات", "GPS بعد استكمال العناوين والإحداثيات"),
        ("التحصيلات والأرصدة", "الصلاحيات المؤسسية بعد تعريف المستخدمين"),
        ("واجهات الإدارة والمندوب والعميل", "التكاملات الخارجية بعد الحسابات والموافقات"),
    ], [4680, 4680])

    add_heading(doc, "بداية عملية لا مشروعًا ضخمًا", 1)
    add_callout(doc, "الخطوة التالية", "اعتماد قائمة الأسعار، اختيار أول 20 عميلًا، تحديد مسؤول الطلبات والتحصيلات، ثم تشغيل تجريبي لمدة عشرة أيام ومراجعة النتائج.")
    add_body(doc, "نسخة العرض المؤقتة الخاصة بالمالك: https://napri-sales-phase-one-2026.logman-gelany.chatgpt.site", bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_body(doc, "للتواصل أو طلب عرض: [البريد] | [الهاتف / واتساب أعمال]", color=GRAY, align=WD_ALIGN_PARAGRAPH.LEFT)
    path = OUT / "04-الكتيب-التعريفي-لنظام-نبري.docx"
    doc.save(path)
    return path


def main():
    paths = [build_proposal(), build_contract(), build_manual(), build_brochure()]
    print("\n".join(str(p) for p in paths))


if __name__ == "__main__":
    main()
